#!/usr/bin/env python3
"""
Markdown 转 Word 文档转换器
支持 Markdown 语法转换为 DOCX 格式，保留格式和样式
"""

import sys
import re
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownToWord:
    """Markdown 到 Word 转换器"""

    def __init__(self):
        self.doc = Document()

        # 默认字体配置
        self.font_cn = '仿宋_GB2312'
        self.font_en = 'Times New Roman'
        self.title_font = '黑体'
        self.heading_font = '楷体_GB2312'

        # 字号配置（单位：磅）
        self.title_size = 22      # 二号
        self.heading1_size = 16   # 三号
        self.heading2_size = 15   # 小三号
        self.heading3_size = 14   # 四号
        self.body_size = 12       # 小四号

        # 行距（磅）
        self.line_spacing = 28

        # 代码块字体
        self.code_font = 'Consolas'
        self.code_size = 10
        self.code_bg_color = 'F5F5F5'

        # 页面设置
        self._setup_page()

    def _setup_page(self):
        """设置页面格式"""
        for section in self.doc.sections:
            section.page_width = Cm(21)   # A4
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

    def _set_run_font(self, run, font_cn=None, font_en=None, size=None, bold=False, italic=False):
        """设置 run 的字体样式"""
        font_cn = font_cn or self.font_cn
        font_en = font_en or self.font_en
        size = size or self.body_size

        run.font.name = font_en
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic

        # 设置中文字体
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_cn)
        rFonts.set(qn('w:ascii'), font_en)
        rFonts.set(qn('w:hAnsi'), font_en)

    def _parse_inline_formatting(self, text):
        """解析行内格式（加粗、斜体、代码等）"""
        # 按换行分割，逐行处理
        lines = text.split('\n')
        for line in lines:
            if not line.strip():
                continue
            self._add_formatted_paragraph(line)

    def _add_formatted_paragraph(self, text):
        """添加带有行内格式的段落 - 优化版本"""
        para = self.doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = Pt(self.line_spacing)

        # 使用 re.split 一次性分割所有标记，更高效
        # 匹配所有行内格式标记
        pattern = r'(`[^`]+`|\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|~~.+?~~|\[[^\]]+\]\([^)]+\))'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            # 代码片段：`code`
            code_match = re.match(r'^`([^`]+)`$', part)
            if code_match:
                run = para.add_run(code_match.group(1))
                self._set_run_font(run, font_cn=self.code_font, font_en=self.code_font,
                                   size=self.code_size)
                continue

            # 加粗：**text** 或 __text__
            bold_match = re.match(r'^\*\*(.+?)\*\*$', part)
            if not bold_match:
                bold_match = re.match(r'^__(.+?)__$', part)
            if bold_match:
                run = para.add_run(bold_match.group(1))
                self._set_run_font(run, bold=True)
                continue

            # 斜体：*text* 或 _text_
            italic_match = re.match(r'^\*(.+?)\*$', part)
            if not italic_match:
                italic_match = re.match(r'^_(.+?)_$', part)
            if italic_match:
                run = para.add_run(italic_match.group(1))
                self._set_run_font(run, italic=True)
                continue

            # 删除线：~~text~~
            strike_match = re.match(r'^~~(.+?)~~$', part)
            if strike_match:
                run = para.add_run(strike_match.group(1))
                self._set_run_font(run)
                run.font.strike = True
                continue

            # 链接：[text](url)
            link_match = re.match(r'^\[([^\]]+)\]\(([^)]+)\)$', part)
            if link_match:
                run = para.add_run(link_match.group(1))
                self._set_run_font(run)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.font.underline = True
                continue

            # 普通文本
            run = para.add_run(part)
            self._set_run_font(run)

    def _add_heading(self, text, level):
        """添加标题"""
        para = self.doc.add_paragraph()
        text = text.strip()

        # 标题字体和大小
        if level == 1:
            size = self.title_size
            font = self.title_font
            align = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            size = self.heading1_size
            font = self.heading_font
            align = WD_ALIGN_PARAGRAPH.LEFT
        elif level == 3:
            size = self.heading2_size
            font = self.heading_font
            align = WD_ALIGN_PARAGRAPH.LEFT
        else:
            size = self.heading3_size
            font = self.font_cn
            align = WD_ALIGN_PARAGRAPH.LEFT

        run = para.add_run(text)
        self._set_run_font(run, font_cn=font, size=size, bold=True)
        para.paragraph_format.alignment = align
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    def _add_code_block(self, code, lang=None):
        """添加代码块"""
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Twips(72)  # 左缩进
        para.paragraph_format.right_indent = Twips(72)  # 右缩进
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        run = para.add_run(code)
        self._set_run_font(run, font_cn=self.code_font, font_en=self.code_font,
                           size=self.code_size)

        # 添加灰色背景
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), self.code_bg_color)
        pPr.append(shd)

    def _add_list_item(self, text, level=0, ordered=False):
        """添加列表项"""
        para = self.doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = Pt(self.line_spacing)
        para.paragraph_format.left_indent = Twips(720 * level)  # 每级缩进

        # 列表标记
        if ordered:
            prefix = f"{level + 1}. "
        else:
            prefix = "• "

        run = para.add_run(f"{prefix}{text}")
        self._set_run_font(run)

    def _add_blockquote(self, text):
        """添加引用块"""
        para = self.doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = Pt(self.line_spacing)
        para.paragraph_format.left_indent = Twips(360)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        # 添加左侧边框模拟引用线
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '808080')
        pBdr.append(left)
        pPr.append(pBdr)

        run = para.add_run(text)
        self._set_run_font(run, italic=True)

    def _add_table(self, rows, headers=None):
        """添加表格"""
        if not rows:
            return

        table = self.doc.add_table(rows=len(rows) if not headers else len(rows) + 1,
                                   cols=len(rows[0]) if rows else 0)
        table.style = 'Table Grid'

        # 设置表头
        if headers:
            header_row = table.rows[0]
            for i, header_text in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header_text
                for para in cell.paragraphs:
                    for run in para.runs:
                        self._set_run_font(run, bold=True)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置表内容
        for row_idx, row_data in enumerate(rows):
            table_row = table.rows[row_idx + 1] if headers else table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(table_row.cells):
                    cell = table_row.cells[col_idx]
                    cell.text = str(cell_text)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            self._set_run_font(run)

    def _parse_markdown(self, md_content):
        """解析 Markdown 内容"""
        lines = md_content.split('\n')
        i = 0
        in_code_block = False
        code_block_content = []
        code_block_lang = None

        while i < len(lines):
            line = lines[i]

            # 代码块开始
            code_block_start = re.match(r'^```(\w*)', line)
            if code_block_start and not in_code_block:
                in_code_block = True
                code_block_lang = code_block_start.group(1)
                code_block_content = []
                i += 1
                continue

            # 代码块结束
            if in_code_block and line.strip() == '```':
                in_code_block = False
                self._add_code_block('\n'.join(code_block_content), code_block_lang)
                code_block_content = []
                code_block_lang = None
                i += 1
                continue

            # 代码块内容
            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue

            # 跳过空行
            if not line.strip():
                i += 1
                continue

            # 标题
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self._add_heading(text, level)
                i += 1
                continue

            # 无序列表
            unordered_match = re.match(r'^[\-\*\+]\s+(.+)$', line)
            if unordered_match:
                self._add_list_item(unordered_match.group(1), level=0, ordered=False)
                i += 1
                continue

            # 有序列表
            ordered_match = re.match(r'^(\d+)\.\s+(.+)$', line)
            if ordered_match:
                self._add_list_item(ordered_match.group(2), level=0, ordered=True)
                i += 1
                continue

            # 引用块
            blockquote_match = re.match(r'^>\s*(.+)$', line)
            if blockquote_match:
                self._add_blockquote(blockquote_match.group(1))
                i += 1
                continue

            # 表格（简单处理）
            if re.match(r'^\|', line) and not re.match(r'^\|\s*[-:|]+\s*\|$', line):
                # 跳过表格分隔线
                if re.match(r'^\|\s*[-:|]+\s*\|$', lines[i + 1] if i + 1 < len(lines) else ''):
                    i += 2
                    continue
                # 解析表格行
                table_rows = []
                while i < len(lines) and lines[i].startswith('|'):
                    if re.match(r'^\|\s*[-:|]+\s*\|$', lines[i]):
                        i += 1
                        continue
                    cells = [c.strip() for c in lines[i].split('|')[1:-1]]
                    table_rows.append(cells)
                    i += 1
                if table_rows:
                    headers = table_rows[0] if len(table_rows) > 1 else None
                    data = table_rows[1:] if len(table_rows) > 1 else table_rows
                    self._add_table(data, headers)
                continue

            # 普通段落
            self._add_formatted_paragraph(line)
            i += 1

    def convert(self, input_path, output_path=None):
        """转换 Markdown 文件到 Word"""
        input_path = Path(input_path)
        if not input_path.exists():
            raise FileNotFoundError(f"文件不存在：{input_path}")

        if output_path is None:
            output_path = input_path.with_suffix('.docx')
        else:
            output_path = Path(output_path)

        # 读取 Markdown 内容
        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # 解析并转换
        self._parse_markdown(md_content)

        # 保存
        self.doc.save(output_path)

        return str(output_path)


def main():
    parser = argparse.ArgumentParser(description='Markdown 转 Word 文档转换器')
    parser.add_argument('input', help='输入的 Markdown 文件路径')
    parser.add_argument('output', nargs='?', help='输出的 Word 文件路径（可选，默认为同名.docx）')
    parser.add_argument('--preset', choices=['default', 'official', 'academic'], default='default',
                        help='预设格式：default（默认）、official（公文）、academic（学术）')

    args = parser.parse_args()

    converter = MarkdownToWord()

    # 应用预设
    if args.preset == 'official':
        converter.font_cn = '仿宋_GB2312'
        converter.title_font = '方正小标宋简体'
        converter.heading_font = '黑体'
        converter.title_size = 22
        converter.heading1_size = 16
        converter.body_size = 16
    elif args.preset == 'academic':
        converter.font_cn = '宋体'
        converter.title_font = '黑体'
        converter.heading_font = '黑体'
        converter.title_size = 18
        converter.heading1_size = 15
        converter.body_size = 12

    try:
        output_file = converter.convert(args.input, args.output)
        print(f'转换成功：{output_file}')
    except Exception as e:
        print(f'转换失败：{e}', file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
