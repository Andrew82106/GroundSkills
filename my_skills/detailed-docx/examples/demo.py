"""
demo.py
-------
用法演示：
  1. 创建一个包含丰富格式的测试文档
  2. 用 DocxEditor 读取结构、替换文本、修改格式
  3. 保存后打开可在 Word 中验证效果
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from my_docx import DocxEditor


def create_demo_document(path: str):
    """创建一个格式丰富的演示文档。"""
    doc = Document()

    # 标题
    h = doc.add_heading('2023年度工作报告', level=1)

    # 正文段落（多格式混排）
    p1 = doc.add_paragraph()
    r1 = p1.add_run('北京公司')
    r1.font.bold = True
    r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(0x00, 0x51, 0x8A)  # 深蓝
    r2 = p1.add_run('在2023年取得了')
    r2.font.size = Pt(12)
    r3 = p1.add_run('非常优秀')
    r3.font.bold = True
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色
    r4 = p1.add_run('的业绩，利润同比增长25%。')
    r4.font.size = Pt(12)

    # 第二段
    p2 = doc.add_paragraph('重点关注条款：本年度合同执行率达到95%，超额完成目标。')
    p2.paragraph_format.space_before = Pt(12)
    p2.paragraph_format.space_after = Pt(6)

    # 表格
    doc.add_heading('业绩汇总表', level=2)
    table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
    headers = ['部门', '目标（万元）', '实际（万元）']
    for i, h_text in enumerate(headers):
        table.cell(0, i).text = h_text
    data = [
        ['销售部', '500', '620'],
        ['技术部', '300', '285'],
        ['运营部', '200', '210'],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            table.cell(r + 1, c).text = val

    doc.save(path)
    print(f"✅ 演示文档已创建: {path}")


def run_demo():
    demo_dir = os.path.dirname(os.path.abspath(__file__))
    input_path = os.path.join(demo_dir, 'demo_input.docx')
    output_path = os.path.join(demo_dir, 'demo_output.docx')

    # Step 1: 创建演示文档
    create_demo_document(input_path)

    # Step 2: 加载文档
    editor = DocxEditor(input_path)

    # Step 3: 查看结构地图
    print("\n📄 文档结构地图：")
    smap = editor.get_structural_map()
    print(f"  段落数: {len(smap['paragraphs'])}")
    for p in smap['paragraphs']:
        print(f"  [{p['index']}] ({p['style']}) {p['text'][:50]}...")
    print(f"  表格数: {len(smap['tables'])}")
    for t in smap['tables']:
        print(f"  [表格{t['index']}] {t['rows']}行 × {t['cols']}列")

    # Step 4: 查看 Run 级别细节
    print("\n🔍 第1段 Run 细节：")
    details = editor.get_run_details(1)
    for d in details:
        print(f"  Run[{d['run_index']}]: \"{d['text']}\" → {d['format']}")

    # Step 5: 替换文本（保留原格式）
    print("\n🔄 替换 '北京公司' → '上海分公司'：")
    result = editor.replace_text('北京公司', '上海分公司')
    print(f"  结果: {result}")

    # Step 6: 替换表格内容
    print("\n🔄 替换表格中 '285' → '310'：")
    result = editor.replace_in_table(0, 2, 2, '285', '310')
    print(f"  结果: {result}")

    # Step 7: 增量格式修改（给 '非常优秀' 加下划线，不影响原有红色粗体斜体）
    print("\n🎨 给 '非常优秀' 加下划线（保留原红色粗体斜体）：")
    result = editor.modify_format('非常优秀', {'underline': True})
    print(f"  结果: {result}")

    # Step 8: 给 '重点关注条款' 加高亮黄色背景
    print("\n🎨 给 '重点关注条款' 加粗体 + 蓝色：")
    result = editor.modify_format('重点关注条款', {'bold': True, 'color': '0000FF'})
    print(f"  结果: {result}")

    # Step 9: 修改段落对齐
    print("\n📐 将标题居中：")
    result = editor.modify_paragraph_format(0, {'alignment': 'center'})
    print(f"  结果: {result}")

    # Step 10: 设置表格单元格背景色
    print("\n🎨 设置表头背景色为浅蓝色：")
    for col in range(3):
        editor.modify_table_cell_shading(0, 0, col, '4472C4')

    # Step 11: 验证修改结果
    print("\n✅ 修改后 Run 细节：")
    details = editor.get_run_details(1)
    for d in details:
        print(f"  Run[{d['run_index']}]: \"{d['text']}\" → {d['format']}")

    # Step 12: 保存
    saved = editor.save(output_path)
    print(f"\n💾 已保存到: {saved}")
    print("请用 Microsoft Word 打开查看效果！")


if __name__ == '__main__':
    run_demo()
