"""
editor.py
---------
DocxEditor — 面向用户和 AI Agent 的核心 API 类。

整合 traverser、run_ops、style_ops 三个模块的能力，
对外暴露直观的文档读取与修改接口。

设计原则：
  - 所有读取方法返回 Python 原生 dict/list，便于 Agent 解析和代码处理。
  - 所有修改方法返回操作结果摘要（替换了多少处等），便于 Agent 做下一步决策。
  - 增量格式修改（方案 A）：仅叠加指定属性，不清除原有格式。
"""

from typing import Dict, Any, List, Optional, Union

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Pt, Inches, Cm, Emu, RGBColor, Length
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .traverser import DocumentTraverser, LocatedParagraph
from .run_ops import replace_text_in_paragraph, apply_format_to_paragraph_text, _build_run_map, _is_text_run
from .style_ops import (
    read_font, write_font, read_run,
    read_paragraph_format, write_paragraph_format,
    read_cell_shading, set_cell_shading, set_cell_border,
)


class DocxEditor:
    """
    Word 文档精细编辑器。

    支持三大类操作：
      1. 从零创建：创建新文档、添加段落/标题/表格
      2. 编辑修改：替换文本（保留格式）、增量修改格式、修改表格
      3. 删除操作：删除段落、表格、表格行、清空单元格

    用法（编辑已有文档）：
        >>> editor = DocxEditor("input.docx")
        >>> editor.get_structural_map()
        >>> editor.replace_text("旧词", "新词")
        >>> editor.save("output.docx")

    用法（从零创建文档）：
        >>> editor = DocxEditor.create_new()
        >>> editor.add_heading("标题", level=1)
        >>> editor.add_paragraph("正文内容", fmt={"size": 12})
        >>> editor.add_table(3, 4, data=[["A", "B", "C", "D"], ...])
        >>> editor.save("new_doc.docx")
    """

    def __init__(self, file_path: str):
        """
        加载一个已有的 .docx 文件。

        Args:
            file_path: .docx 文件路径
        """
        self._path = file_path
        self._doc = Document(file_path)
        self._traverser = DocumentTraverser(self._doc)

    @classmethod
    def create_new(cls, template_path: Optional[str] = None) -> 'DocxEditor':
        """
        创建一个新的空白文档（或基于模板创建）。

        Args:
            template_path: 可选的模板文件路径。None = 使用 python-docx 默认空白模板。

        Returns:
            DocxEditor 实例（未关联文件路径，save 时必须指定 output_path）
        """
        instance = object.__new__(cls)
        instance._path = None
        if template_path:
            instance._doc = Document(template_path)
        else:
            instance._doc = Document()
        instance._traverser = DocumentTraverser(instance._doc)
        return instance

    # ═══════════════════════════════════════════════
    #  读取类 API
    # ═══════════════════════════════════════════════

    def get_structural_map(self) -> Dict[str, Any]:
        """
        返回文档的完整结构地图（JSON 友好的 dict）。

        返回：
        {
            "paragraphs": [
                {"index": 0, "style": "Heading 1", "text": "第一章 概述", "alignment": "left"},
                {"index": 1, "style": "Normal", "text": "这是正文内容...", "alignment": null},
                ...
            ],
            "tables": [
                {
                    "index": 0, "rows": 3, "cols": 4,
                    "preview": [["姓名", "年龄", ...], ["张三", "25", ...], ...]
                },
                ...
            ],
            "sections": [
                {"index": 0, "has_header": True, "has_footer": True,
                 "header_text": "公司名称", "footer_text": "第1页"},
                ...
            ]
        }
        """
        result: Dict[str, Any] = {}

        # 段落
        paras = []
        for idx, para in enumerate(self._doc.paragraphs):
            paras.append({
                'index': idx,
                'style': para.style.name if para.style else None,
                'text': para.text,
                'alignment': self._alignment_str(para.alignment),
            })
        result['paragraphs'] = paras

        # 表格
        tables = []
        for t_idx, table in enumerate(self._doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            preview = []
            for row in table.rows:
                preview.append([cell.text for cell in row.cells])
            tables.append({
                'index': t_idx,
                'rows': rows,
                'cols': cols,
                'preview': preview,
            })
        result['tables'] = tables

        # 节（Sections）
        sections = []
        for s_idx, section in enumerate(self._doc.sections):
            header = section.header
            footer = section.footer
            # 读取分栏信息
            col_info = self._read_cols_from_section(section)
            sections.append({
                'index': s_idx,
                'has_header': not header.is_linked_to_previous,
                'has_footer': not footer.is_linked_to_previous,
                'header_text': '\n'.join(p.text for p in header.paragraphs) if not header.is_linked_to_previous else '',
                'footer_text': '\n'.join(p.text for p in footer.paragraphs) if not footer.is_linked_to_previous else '',
                'columns': col_info,
            })
        result['sections'] = sections

        return result

    def get_paragraphs(self, scope: str = 'body') -> List[Dict[str, Any]]:
        """
        获取段落的详细信息列表。

        Args:
            scope: 'body' | 'all' | 'tables' | 'headers' | 'footers'

        Returns:
            列表，每个元素：
            {
                "location": {...},
                "text": "段落文本",
                "style": "Normal",
                "paragraph_format": {...},
                "runs": [{"text": "...", "format": {...}}, ...]
            }
        """
        result = []
        for lp in self._traverser.iter_all(scope):
            para = lp.paragraph
            runs_info = []
            for run in para.runs:
                runs_info.append(read_run(run))

            result.append({
                'location': lp.location.to_dict(),
                'text': para.text,
                'style': para.style.name if para.style else None,
                'paragraph_format': read_paragraph_format(para),
                'runs': runs_info,
            })
        return result

    def get_table_data(self, table_idx: int) -> Dict[str, Any]:
        """
        获取指定表格的完整数据（含每个单元格的文本和格式）。

        Args:
            table_idx: 表格索引（从 0 开始）

        Returns:
            {
                "rows": 3,
                "cols": 4,
                "cells": [
                    [  // row 0
                        {"text": "...", "shading": "FFFFFF", "paragraphs": [...]},
                        ...
                    ],
                    ...
                ]
            }
        """
        tables = self._doc.tables
        if table_idx >= len(tables):
            raise IndexError(f"表格索引 {table_idx} 超出范围，文档中共有 {len(tables)} 个表格")

        table = tables[table_idx]
        cells_data = []
        for r_idx, row in enumerate(table.rows):
            row_data = []
            for c_idx, cell in enumerate(row.cells):
                paras_info = []
                for para in cell.paragraphs:
                    paras_info.append({
                        'text': para.text,
                        'paragraph_format': read_paragraph_format(para),
                        'runs': [read_run(run) for run in para.runs],
                    })
                row_data.append({
                    'text': cell.text,
                    'shading': read_cell_shading(cell),
                    'paragraphs': paras_info,
                })
            cells_data.append(row_data)

        return {
            'rows': len(table.rows),
            'cols': len(table.columns),
            'cells': cells_data,
        }

    def get_table_cell_text(self, table_idx: int, row_idx: int, col_idx: int) -> str:
        """获取指定表格单元格的文本。"""
        return self._get_cell(table_idx, row_idx, col_idx).text

    def get_run_details(self, para_idx: int) -> List[Dict[str, Any]]:
        """
        获取正文中第 para_idx 段的 Run 级别详细信息。
        对调试格式问题非常有用：可以看清楚每个 Run 的文本碎片和精确格式。
        """
        paras = self._doc.paragraphs
        if para_idx >= len(paras):
            raise IndexError(f"段落索引 {para_idx} 超出范围，正文共 {len(paras)} 段")
        para = paras[para_idx]
        details = []
        for r_idx, run in enumerate(para.runs):
            details.append({
                'run_index': r_idx,
                'text': run.text,
                'is_text_run': _is_text_run(run),
                'format': read_font(run.font),
            })
        return details

    # ═══════════════════════════════════════════════
    #  文本替换 API
    # ═══════════════════════════════════════════════

    def replace_text(
        self,
        old: str,
        new: str,
        scope: str = 'all',
        count: int = 0,
        use_regex: bool = False,
    ) -> Dict[str, Any]:
        """
        全局替换文本，自动保留第一个命中 Run 的格式。

        Args:
            old: 旧文本（或正则表达式）
            new: 替换后的文本
            scope: 搜索范围 'all' | 'body' | 'tables' | 'headers' | 'footers'
            count: 最多替换次数，0 = 全部
            use_regex: 是否将 old 作为正则表达式

        Returns:
            {"replaced_count": 5, "scope": "all"}
        """
        total = 0
        remaining = count
        for lp in self._traverser.iter_all(scope):
            per_para_count = remaining if remaining > 0 else 0
            n = replace_text_in_paragraph(lp.paragraph, old, new, count=per_para_count, use_regex=use_regex)
            total += n
            if count > 0:
                remaining -= n
                if remaining <= 0:
                    break
        return {'replaced_count': total, 'scope': scope}

    def replace_in_paragraph(
        self,
        para_idx: int,
        old: str,
        new: str,
        count: int = 0,
        use_regex: bool = False,
    ) -> Dict[str, Any]:
        """
        替换正文中指定段落内的文本。

        Args:
            para_idx: 段落索引
            old, new, count, use_regex: 同 replace_text

        Returns:
            {"replaced_count": N, "para_idx": para_idx}
        """
        paras = self._doc.paragraphs
        if para_idx >= len(paras):
            raise IndexError(f"段落索引 {para_idx} 超出范围，正文共 {len(paras)} 段")
        n = replace_text_in_paragraph(paras[para_idx], old, new, count=count, use_regex=use_regex)
        return {'replaced_count': n, 'para_idx': para_idx}

    def replace_in_table(
        self,
        table_idx: int,
        row_idx: int,
        col_idx: int,
        old: str,
        new: str,
        count: int = 0,
        use_regex: bool = False,
    ) -> Dict[str, Any]:
        """
        替换指定表格单元格内的文本。

        Returns:
            {"replaced_count": N, "table_idx": ..., "row_idx": ..., "col_idx": ...}
        """
        cell = self._get_cell(table_idx, row_idx, col_idx)
        total = 0
        for para in cell.paragraphs:
            n = replace_text_in_paragraph(para, old, new, count=count, use_regex=use_regex)
            total += n
            if count > 0:
                count -= n
                if count <= 0:
                    break
        return {
            'replaced_count': total,
            'table_idx': table_idx, 'row_idx': row_idx, 'col_idx': col_idx,
        }

    # ═══════════════════════════════════════════════
    #  格式修改 API（增量式叠加 - 方案 A）
    # ═══════════════════════════════════════════════

    def modify_format(
        self,
        target_text: str,
        fmt: Dict[str, Any],
        scope: str = 'all',
        count: int = 0,
        use_regex: bool = False,
    ) -> Dict[str, Any]:
        """
        查找 target_text，对命中的文本区间做增量格式叠加。

        不会清除文本上已有的其他格式属性，仅写入 fmt 中指定的属性。

        示例：
            # 原文 "项目非常成功" 中 "非常" 是红色粗体
            editor.modify_format("项目非常成功", {"underline": True})
            # 结果："非常" 变成 红色+粗体+下划线，其余字变成 下划线

        Args:
            target_text: 要查找的文本或 regex
            fmt: 格式字典，如 {'bold': True, 'color': 'FF0000', 'size': 14}
            scope: 搜索范围
            count: 最多处理几处
            use_regex: 是否正则

        Returns:
            {"modified_count": 3, "scope": "all"}
        """
        total = 0
        remaining = count
        for lp in self._traverser.iter_all(scope):
            per_para_count = remaining if remaining > 0 else 0
            n = apply_format_to_paragraph_text(
                lp.paragraph, target_text, fmt,
                count=per_para_count, use_regex=use_regex,
            )
            total += n
            if count > 0:
                remaining -= n
                if remaining <= 0:
                    break
        return {'modified_count': total, 'scope': scope}

    def modify_paragraph_format(
        self,
        para_idx: int,
        fmt: Dict[str, Any],
    ) -> Dict[str, Any]:
        """
        修改正文中第 para_idx 段的段落级格式（对齐、缩进、行距等）。

        Args:
            para_idx: 段落索引
            fmt: 格式字典，如 {'alignment': 'center', 'space_before': 12}

        Returns:
            {"success": True, "para_idx": para_idx}
        """
        paras = self._doc.paragraphs
        if para_idx >= len(paras):
            raise IndexError(f"段落索引 {para_idx} 超出范围，正文共 {len(paras)} 段")
        write_paragraph_format(paras[para_idx], fmt)
        return {'success': True, 'para_idx': para_idx}

    def modify_table_cell_shading(
        self,
        table_idx: int,
        row_idx: int,
        col_idx: int,
        hex_color: str,
    ) -> Dict[str, Any]:
        """设置表格单元格的背景色。"""
        cell = self._get_cell(table_idx, row_idx, col_idx)
        set_cell_shading(cell, hex_color)
        return {
            'success': True,
            'table_idx': table_idx, 'row_idx': row_idx, 'col_idx': col_idx,
            'color': hex_color,
        }

    def modify_table_cell_border(
        self,
        table_idx: int,
        row_idx: int,
        col_idx: int,
        **edges,
    ) -> Dict[str, Any]:
        """
        设置表格单元格边框。

        使用方式：
            editor.modify_table_cell_border(0, 0, 0,
                top={'sz': 12, 'val': 'single', 'color': '000000'},
                bottom={'sz': 12, 'val': 'single', 'color': '000000'})
        """
        cell = self._get_cell(table_idx, row_idx, col_idx)
        set_cell_border(cell, **edges)
        return {
            'success': True,
            'table_idx': table_idx, 'row_idx': row_idx, 'col_idx': col_idx,
        }

    def apply_three_line_style(
        self,
        table_idx: int,
        header_rows: int = 1,
        top_border_pt: float = 1.5,
        mid_border_pt: float = 0.75,
        bottom_border_pt: float = 1.5,
        color: str = '000000',
    ) -> Dict[str, Any]:
        """
        将表格设置为三线表（booktabs）样式 —— 中文学术论文的标准表格格式。

        三线表只保留三条横线：
        - 顶线（top rule）：粗线，表格最上方
        - 栏目线（mid rule）：细线，分隔表头与表体
        - 底线（bottom rule）：粗线，表格最下方
        - 所有竖线和其他横线均清除

        Args:
            table_idx: 表格索引
            header_rows: 表头行数（默认1），栏目线在第 header_rows 行之后
            top_border_pt: 顶线粗细（磅值），默认 1.5
            mid_border_pt: 栏目线粗细（磅值），默认 0.75
            bottom_border_pt: 底线粗细（磅值），默认 1.5
            color: 线条颜色（hex），默认黑色

        Returns:
            {"success": True, "table_idx": N, "header_rows": M}
        """
        tables = self._doc.tables
        if table_idx >= len(tables):
            raise IndexError(f"表格索引 {table_idx} 超出范围，共 {len(tables)} 个表格")
        table = tables[table_idx]

        num_rows = len(table.rows)
        if header_rows >= num_rows:
            raise ValueError(f"header_rows ({header_rows}) 必须小于表格总行数 ({num_rows})")

        # 尺寸转换: pt → 1/8 pt (Word 的 w:sz 单位)
        top_sz = str(int(top_border_pt * 8))
        mid_sz = str(int(mid_border_pt * 8))
        bot_sz = str(int(bottom_border_pt * 8))

        no_border = {'val': 'nil'}
        top_rule = {'sz': top_sz, 'val': 'single', 'color': color, 'space': '0'}
        mid_rule = {'sz': mid_sz, 'val': 'single', 'color': color, 'space': '0'}
        bot_rule = {'sz': bot_sz, 'val': 'single', 'color': color, 'space': '0'}

        # 第1步：清除表格级边框（tblBorders）并设置为无边框
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # 移除已有的 tblBorders
        for old in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old)

        tblBorders = OxmlElement('w:tblBorders')
        for edge_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge_name}')
            for k, v in no_border.items():
                el.set(qn(f'w:{k}'), v)
            tblBorders.append(el)
        tblPr.append(tblBorders)

        # 第2步：清除所有单元格级边框
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    for old in tcPr.findall(qn('w:tcBorders')):
                        tcPr.remove(old)

        # 第3步：设置三线
        # 为每一行累积需要设置的边框，最后一次性设置，避免 set_cell_border 的覆盖
        # row_borders[row_idx] = dict of edge_name → rule
        row_borders = {}

        # 顶线：第一行的 top
        row_borders.setdefault(0, {})['top'] = top_rule

        # 栏目线：header 最后一行的 bottom
        hr = header_rows - 1
        row_borders.setdefault(hr, {})['bottom'] = mid_rule

        # 底线：最后一行的 bottom
        last = num_rows - 1
        row_borders.setdefault(last, {})['bottom'] = bot_rule

        for row_idx, edges in row_borders.items():
            for cell in table.rows[row_idx].cells:
                set_cell_border(cell, **edges)

        return {'success': True, 'table_idx': table_idx, 'header_rows': header_rows}


    # ═══════════════════════════════════════════════
    #  保存
    # ═══════════════════════════════════════════════

    def save(self, output_path: Optional[str] = None) -> str:
        """
        保存文档。

        Args:
            output_path: 输出路径。None 则覆盖原文件。

        Returns:
            保存路径字符串
        """
        path = output_path or self._path
        if path is None:
            raise ValueError("未指定保存路径。使用 create_new() 创建的文档必须指定 output_path。")
        self._doc.save(path)
        return path

    # ═══════════════════════════════════════════════
    #  创建类 API（从零构建文档）
    # ═══════════════════════════════════════════════

    def add_heading(self, text: str, level: int = 1) -> Dict[str, Any]:
        """
        添加一个标题段落。

        Args:
            text: 标题文字
            level: 标题级别 (0=Title, 1=Heading1, 2=Heading2, ...)

        Returns:
            {"success": True, "para_idx": N}
        """
        self._doc.add_heading(text, level=level)
        idx = len(self._doc.paragraphs) - 1
        return {'success': True, 'para_idx': idx}

    def add_paragraph(
        self,
        text: str = '',
        style: Optional[str] = None,
        fmt: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        添加一个新段落。

        Args:
            text: 段落文字
            style: 段落样式名（如 'Normal', 'List Bullet'）
            fmt: 字体格式字典，应用于整个段落文本
                 如 {'bold': True, 'size': 12, 'color': 'FF0000'}

        Returns:
            {"success": True, "para_idx": N}
        """
        para = self._doc.add_paragraph(text, style=style)
        if fmt and para.runs:
            from .run_ops import _apply_fmt_to_run
            _apply_fmt_to_run(para.runs[0], fmt)
        idx = len(self._doc.paragraphs) - 1
        return {'success': True, 'para_idx': idx}

    def add_run_to_paragraph(
        self,
        para_idx: int,
        text: str,
        fmt: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        在已有段落末尾追加一个 Run（可设置独立格式）。

        用途：在同一段落内创建多格式混排文本。

        Args:
            para_idx: 目标段落索引
            text: 追加的文本
            fmt: 该 Run 的字体格式，如 {'bold': True, 'color': 'FF0000'}

        Returns:
            {"success": True, "para_idx": N, "run_idx": M}
        """
        paras = self._doc.paragraphs
        if para_idx >= len(paras):
            raise IndexError(f"段落索引 {para_idx} 超出范围，正文共 {len(paras)} 段")
        run = paras[para_idx].add_run(text)
        if fmt:
            from .run_ops import _apply_fmt_to_run
            _apply_fmt_to_run(run, fmt)
        return {'success': True, 'para_idx': para_idx, 'run_idx': len(paras[para_idx].runs) - 1}

    def add_table(
        self,
        rows: int,
        cols: int,
        data: Optional[List[List[str]]] = None,
        style: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        添加一个新表格。

        Args:
            rows: 行数
            cols: 列数
            data: 可选的二维字符串列表填充数据，如 [["A", "B"], ["C", "D"]]
            style: 表格样式名（如 'Table Grid', 'Light Grid Accent 1'）

        Returns:
            {"success": True, "table_idx": N}
        """
        table = self._doc.add_table(rows=rows, cols=cols)
        if style:
            try:
                table.style = style
            except KeyError:
                pass  # 样式不存在则忽略
        if data:
            for r_idx, row_data in enumerate(data):
                if r_idx >= rows:
                    break
                for c_idx, val in enumerate(row_data):
                    if c_idx >= cols:
                        break
                    table.cell(r_idx, c_idx).text = str(val)
        idx = len(self._doc.tables) - 1
        return {'success': True, 'table_idx': idx}

    # ═══════════════════════════════════════════════
    #  图片 API（插入 / 统计）
    # ═══════════════════════════════════════════════

    def add_picture(
        self,
        image_path: str,
        width: Optional[Union[float, Length]] = None,
        height: Optional[Union[float, Length]] = None,
        alignment: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        在文档末尾插入一张图片（作为新段落）。

        Args:
            image_path: 图片文件路径（支持 PNG/JPG/GIF/BMP/TIFF 等）。
            width:  图片宽度。float 视为英寸；可传 Inches(2)/Cm(5)/Emu(...) 等 Length 对象。
                    width/height 都为 None 时使用图片原始尺寸；只指定一边则按比例缩放。
            height: 图片高度，规则同 width。
            alignment: 图片所在段落对齐方式 'left' | 'center' | 'right'，None = 默认（左）。

        Returns:
            {"success": True, "para_idx": N, "image_path": "..."}
        """
        para = self._doc.add_paragraph()
        run = para.add_run()
        run.add_picture(
            image_path,
            width=self._to_length(width),
            height=self._to_length(height),
        )
        if alignment is not None:
            para.alignment = self._alignment_enum(alignment)
        idx = len(self._doc.paragraphs) - 1
        return {'success': True, 'para_idx': idx, 'image_path': image_path}

    def add_picture_to_paragraph(
        self,
        para_idx: int,
        image_path: str,
        width: Optional[Union[float, Length]] = None,
        height: Optional[Union[float, Length]] = None,
    ) -> Dict[str, Any]:
        """
        在正文中已存在的段落末尾追加一张图片（作为新的 Run）。
        适用于"图文混排"场景，例如行内 emoji / 小图标 / 与文字同段的插图。

        Args:
            para_idx: 目标段落索引
            image_path: 图片路径
            width, height: 同 add_picture

        Returns:
            {"success": True, "para_idx": N, "run_idx": M, "image_path": "..."}
        """
        paras = self._doc.paragraphs
        if para_idx >= len(paras):
            raise IndexError(f"段落索引 {para_idx} 超出范围，正文共 {len(paras)} 段")
        run = paras[para_idx].add_run()
        run.add_picture(
            image_path,
            width=self._to_length(width),
            height=self._to_length(height),
        )
        run_idx = len(paras[para_idx].runs) - 1
        return {
            'success': True,
            'para_idx': para_idx,
            'run_idx': run_idx,
            'image_path': image_path,
        }

    def add_picture_to_table_cell(
        self,
        table_idx: int,
        row_idx: int,
        col_idx: int,
        image_path: str,
        width: Optional[Union[float, Length]] = None,
        height: Optional[Union[float, Length]] = None,
        alignment: Optional[str] = None,
        clear_cell: bool = False,
    ) -> Dict[str, Any]:
        """
        在表格单元格中插入图片。

        Args:
            table_idx, row_idx, col_idx: 单元格坐标（支持合并单元格）
            image_path: 图片路径
            width, height: 同 add_picture
            alignment: 图片所在段落的对齐 'left' | 'center' | 'right'
            clear_cell: True = 先清空单元格已有文本/图片再插入；False = 在原内容后追加

        Returns:
            {"success": True, "table_idx": ..., "row_idx": ..., "col_idx": ..., "image_path": "..."}
        """
        cell = self._get_cell(table_idx, row_idx, col_idx)

        if clear_cell:
            for para in cell.paragraphs:
                p_el = para._p
                p_el.getparent().remove(p_el)
            target_para = cell.add_paragraph()
        else:
            if cell.paragraphs and not cell.paragraphs[-1].runs and not cell.paragraphs[-1].text:
                target_para = cell.paragraphs[-1]
            else:
                target_para = cell.add_paragraph()

        run = target_para.add_run()
        run.add_picture(
            image_path,
            width=self._to_length(width),
            height=self._to_length(height),
        )
        if alignment is not None:
            target_para.alignment = self._alignment_enum(alignment)

        return {
            'success': True,
            'table_idx': table_idx,
            'row_idx': row_idx,
            'col_idx': col_idx,
            'image_path': image_path,
        }

    # ═══════════════════════════════════════════════
    #  公式 API（LaTeX → OMML）
    # ═══════════════════════════════════════════════

    def add_equation(
        self,
        latex: str,
        alignment: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        在文档末尾插入 LaTeX 公式（生成新段落）。

        转换链路：LaTeX → MathML（latex2mathml）→ OMML（mathml2omml）→ Word 段落。
        生成的是 Word 原生公式（<m:oMath>），可在 Word 公式编辑器中继续编辑。

        若任意一步转换失败，会在该位置写入文本 "[渲染失败] {原始 LaTeX}"，
        以便人工排查；该方法本身不抛异常。

        Args:
            latex: LaTeX 源代码，例如 r"E = mc^2"、r"\\frac{a}{b}"
            alignment: 段落对齐 'left' | 'center' | 'right'，None = 默认

        Returns:
            {"success": True, "para_idx": N, "rendered": True}
            渲染失败时: {"success": True, "para_idx": N, "rendered": False, "error": "..."}
        """
        para = self._doc.add_paragraph()
        result = self._insert_equation_into_paragraph(para, latex)
        if alignment is not None:
            para.alignment = self._alignment_enum(alignment)
        idx = len(self._doc.paragraphs) - 1
        return {'success': True, 'para_idx': idx, **result}

    def add_equation_to_paragraph(
        self,
        para_idx: int,
        latex: str,
    ) -> Dict[str, Any]:
        """
        在正文中已存在段落的末尾追加公式（行内公式，与文字同段）。

        Args:
            para_idx: 目标段落索引
            latex: LaTeX 源代码

        Returns:
            {"success": True, "para_idx": N, "rendered": True}
            渲染失败时: {"success": True, "para_idx": N, "rendered": False, "error": "..."}
        """
        paras = self._doc.paragraphs
        if para_idx >= len(paras):
            raise IndexError(f"段落索引 {para_idx} 超出范围，正文共 {len(paras)} 段")
        result = self._insert_equation_into_paragraph(paras[para_idx], latex)
        return {'success': True, 'para_idx': para_idx, **result}

    @staticmethod
    def _preprocess_latex(latex: str) -> str:
        """
        修复 latex2mathml 已知的兼容性问题：
        - \\bigl / \\bigr / \\big / \\Bigl / \\Bigr / \\Big / \\bigg / \\Bigg [lrm]?
          这些大小修饰命令 latex2mathml 不识别，会把跟随的 \\{ \\} 留成字面量
          导致 Word 里显示成 \\( \\}。直接剥掉前缀，保留后面的定界符。
        - \\nicefrac{a}{b} → \\frac{a}{b}
        """
        import re as _re
        latex = _re.sub(r'\\(?:big|Big|bigg|Bigg)[lrm]?(?=[\s\\({\[])', '', latex)
        latex = _re.sub(r'\\nicefrac\b', r'\\frac', latex)
        return latex

    @staticmethod
    def _insert_equation_into_paragraph(paragraph: Paragraph, latex: str) -> Dict[str, Any]:
        """
        把 LaTeX 转成 OMML 元素并追加到指定段落 <w:p> 末尾。
        失败时退回写文本，并返回 {"rendered": False, "error": "..."}。
        """
        try:
            import latex2mathml.converter as _l2m
            import mathml2omml as _m2o
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsmap as _nsmap

            mathml = _l2m.convert(DocxEditor._preprocess_latex(latex))
            omml_str = _m2o.convert(mathml)

            # mathml2omml 输出的 <m:oMath> 没有 xmlns 声明，需要加上
            ns_decl = (
                ' xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
                ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            )
            if omml_str.startswith('<m:oMath>'):
                omml_str = '<m:oMath' + ns_decl + '>' + omml_str[len('<m:oMath>'):]
            elif omml_str.startswith('<m:oMath '):
                # 已带属性，避免重复添加
                pass
            else:
                # 不是预期格式
                raise ValueError(f"mathml2omml 输出非 <m:oMath>: {omml_str[:80]}")

            omath_el = parse_xml(omml_str)
            paragraph._p.append(omath_el)
            return {'rendered': True}
        except Exception as e:
            paragraph.add_run(f"[渲染失败] {latex}")
            return {'rendered': False, 'error': f"{type(e).__name__}: {e}"}

    def count_pictures(self) -> Dict[str, Any]:
        """
        统计文档中内联图片（inline shapes）的数量。

        Returns:
            {"count": N}
        """
        return {'count': len(self._doc.inline_shapes)}

    # ═══════════════════════════════════════════════
    #  分栏（Multi-column）API
    # ═══════════════════════════════════════════════

    def get_section_columns(self, section_idx: int = 0) -> Dict[str, Any]:
        """
        获取指定节的分栏设置。

        Args:
            section_idx: 节索引（从 0 开始）

        Returns:
            {
                "num": 2,              # 栏数
                "space": 1.27,         # 栏间距（厘米），仅 equal_width 时有效
                "equal_width": True,   # 是否等宽分栏
                "separator": False,    # 栏间是否有分隔线
                "details": [           # 仅 equal_width=False 时有内容
                    {"width_cm": 7.5, "space_cm": 1.27},
                    {"width_cm": 7.5, "space_cm": 0},
                ]
            }
        """
        section = self._get_section(section_idx)
        return self._read_cols_from_section(section)

    def set_section_columns(
        self,
        section_idx: int = 0,
        num: int = 2,
        space_cm: float = 1.27,
        equal_width: bool = True,
        separator: bool = False,
        col_widths_cm: Optional[List[float]] = None,
    ) -> Dict[str, Any]:
        """
        设置指定节的分栏布局。

        Args:
            section_idx: 节索引
            num: 栏数（1=单栏, 2=双栏, 3=三栏, ...）
            space_cm: 栏间距（厘米），仅 equal_width=True 时生效
            equal_width: 是否等宽分栏。设 False 时必须通过 col_widths_cm 指定每栏宽度
            separator: 是否在栏间显示分隔线
            col_widths_cm: 不等宽分栏时，各栏宽度列表（厘米），
                           长度必须 == num。栏间距自动按 space_cm 填充。
                           例如 [10, 6] 表示左栏 10cm、右栏 6cm。

        Returns:
            {"success": True, "section_idx": N, "num": 2}
        """
        section = self._get_section(section_idx)
        sectPr = section._sectPr

        # 移除已有的 <w:cols>
        for old_cols in sectPr.findall(qn('w:cols')):
            sectPr.remove(old_cols)

        from lxml import etree
        cols_el = etree.SubElement(sectPr, qn('w:cols'))
        cols_el.set(qn('w:num'), str(num))

        if separator:
            cols_el.set(qn('w:sep'), '1')

        if equal_width or col_widths_cm is None:
            cols_el.set(qn('w:equalWidth'), '1')
            # space 以 twips 为单位（1 cm = 567 twips）
            space_twips = int(round(space_cm * 567))
            cols_el.set(qn('w:space'), str(space_twips))
        else:
            if len(col_widths_cm) != num:
                raise ValueError(
                    f"col_widths_cm 长度 ({len(col_widths_cm)}) 必须等于栏数 ({num})"
                )
            cols_el.set(qn('w:equalWidth'), '0')
            space_twips = int(round(space_cm * 567))
            for i, w_cm in enumerate(col_widths_cm):
                col_el = etree.SubElement(cols_el, qn('w:col'))
                col_el.set(qn('w:w'), str(int(round(w_cm * 567))))
                # 最后一栏不需要 space
                if i < num - 1:
                    col_el.set(qn('w:space'), str(space_twips))

        return {'success': True, 'section_idx': section_idx, 'num': num}

    def add_column_break(self, para_idx: Optional[int] = None) -> Dict[str, Any]:
        """
        插入分栏符（Column Break），强制后续内容转入下一栏。

        在双栏/多栏布局中，用分栏符控制内容在哪一栏断开。
        例如双栏模式下，在左栏末尾插入分栏符，后续内容会强制从右栏顶部开始。

        Args:
            para_idx: 在哪个段落**之前**插入分栏符。
                      None = 在文档末尾追加一个含分栏符的新段落。
                      指定索引时，会在该段落的第一个 Run 前插入 <w:br w:type="column"/>。

        Returns:
            {"success": True, "para_idx": N}
        """
        if para_idx is None:
            # 在文档末尾追加一个新段落，包含分栏符
            para = self._doc.add_paragraph()
            run = para.add_run()
            self._add_column_break_to_run(run)
            idx = len(self._doc.paragraphs) - 1
            return {'success': True, 'para_idx': idx}
        else:
            paras = self._doc.paragraphs
            if para_idx >= len(paras):
                raise IndexError(f"段落索引 {para_idx} 超出范围，正文共 {len(paras)} 段")
            para = paras[para_idx]
            # 在段落开头插入一个含分栏符的 Run
            from lxml import etree
            new_run = etree.SubElement(para._p, qn('w:r'))
            br = etree.SubElement(new_run, qn('w:br'))
            br.set(qn('w:type'), 'column')
            # 将新 Run 移到段落最前面（在所有已有 Run 之前）
            para._p.insert(0, new_run)
            return {'success': True, 'para_idx': para_idx}

    # ═══════════════════════════════════════════════
    #  删除类 API
    # ═══════════════════════════════════════════════

    def delete_paragraph(self, para_idx: int) -> Dict[str, Any]:
        """
        删除正文中第 para_idx 段。

        警告：删除后后续段落的索引会前移。

        Returns:
            {"success": True, "deleted_text": "被删除的文本"}
        """
        paras = self._doc.paragraphs
        if para_idx >= len(paras):
            raise IndexError(f"段落索引 {para_idx} 超出范围，正文共 {len(paras)} 段")
        para = paras[para_idx]
        deleted_text = para.text
        p_element = para._p
        p_element.getparent().remove(p_element)
        return {'success': True, 'deleted_text': deleted_text}

    def delete_table(self, table_idx: int) -> Dict[str, Any]:
        """
        删除第 table_idx 个表格。

        Returns:
            {"success": True, "table_idx": table_idx}
        """
        tables = self._doc.tables
        if table_idx >= len(tables):
            raise IndexError(f"表格索引 {table_idx} 超出范围，共 {len(tables)} 个表格")
        tbl_element = tables[table_idx]._tbl
        tbl_element.getparent().remove(tbl_element)
        return {'success': True, 'table_idx': table_idx}

    def delete_table_row(self, table_idx: int, row_idx: int) -> Dict[str, Any]:
        """
        删除指定表格中的一行。

        Returns:
            {"success": True, "table_idx": ..., "row_idx": ...}
        """
        tables = self._doc.tables
        if table_idx >= len(tables):
            raise IndexError(f"表格索引 {table_idx} 超出范围")
        table = tables[table_idx]
        if row_idx >= len(table.rows):
            raise IndexError(f"行索引 {row_idx} 超出范围，共 {len(table.rows)} 行")
        tr = table.rows[row_idx]._tr
        tr.getparent().remove(tr)
        return {'success': True, 'table_idx': table_idx, 'row_idx': row_idx}

    def clear_table_cell(
        self,
        table_idx: int,
        row_idx: int,
        col_idx: int,
    ) -> Dict[str, Any]:
        """
        清空指定单元格的文本内容（保留单元格结构和格式）。

        Returns:
            {"success": True, "cleared_text": "原内容"}
        """
        cell = self._get_cell(table_idx, row_idx, col_idx)
        old_text = cell.text
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ''
        return {
            'success': True,
            'cleared_text': old_text,
            'table_idx': table_idx, 'row_idx': row_idx, 'col_idx': col_idx,
        }

    def set_table_cell_text(
        self,
        table_idx: int,
        row_idx: int,
        col_idx: int,
        text: str,
        fmt: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        设置指定单元格的文本（覆盖写入）。

        适用于空单元格或需要完全重写内容的场景。
        如果单元格已有 Run，则替换第一个 Run 的文本并清空其余 Run；
        如果为空，则新增一个 Run。可选附带字体格式。

        Args:
            table_idx, row_idx, col_idx: 单元格坐标
            text: 要写入的文本
            fmt: 可选的字体格式字典，如 {'name': '黑体', 'size': 12, 'bold': True}

        Returns:
            {"success": True, "old_text": "原内容", ...}
        """
        cell = self._get_cell(table_idx, row_idx, col_idx)
        old_text = cell.text
        para = cell.paragraphs[0]

        if para.runs:
            para.runs[0].text = text
            for run in para.runs[1:]:
                run.text = ''
            target_run = para.runs[0]
        else:
            target_run = para.add_run(text)

        if fmt:
            from .run_ops import _apply_fmt_to_run
            _apply_fmt_to_run(target_run, fmt)

        return {
            'success': True,
            'old_text': old_text,
            'table_idx': table_idx, 'row_idx': row_idx, 'col_idx': col_idx,
        }

    # ═══════════════════════════════════════════════
    #  内部辅助
    # ═══════════════════════════════════════════════

    def _get_section(self, section_idx: int):
        """获取指定节对象，含边界检查。"""
        sections = self._doc.sections
        if section_idx >= len(sections):
            raise IndexError(
                f"节索引 {section_idx} 超出范围，文档共 {len(sections)} 个节"
            )
        return sections[section_idx]

    @staticmethod
    def _read_cols_from_section(section) -> Dict[str, Any]:
        """
        从节的 sectPr 中读取 <w:cols> 分栏信息。
        如果没有 <w:cols>，返回默认的单栏配置。
        """
        sectPr = section._sectPr
        cols_el = sectPr.find(qn('w:cols'))

        if cols_el is None:
            return {
                'num': 1,
                'space': 0,
                'equal_width': True,
                'separator': False,
                'details': [],
            }

        num = int(cols_el.get(qn('w:num'), '1'))
        space_raw = cols_el.get(qn('w:space'), '0')
        space_cm = round(int(space_raw) / 567, 2)
        equal_width = cols_el.get(qn('w:equalWidth'), '1') != '0'
        separator = cols_el.get(qn('w:sep'), '0') == '1'

        details = []
        for col_el in cols_el.findall(qn('w:col')):
            w = col_el.get(qn('w:w'), '0')
            s = col_el.get(qn('w:space'), '0')
            details.append({
                'width_cm': round(int(w) / 567, 2),
                'space_cm': round(int(s) / 567, 2),
            })

        return {
            'num': num,
            'space': space_cm,
            'equal_width': equal_width,
            'separator': separator,
            'details': details,
        }

    @staticmethod
    def _add_column_break_to_run(run):
        """在指定 Run 中插入 <w:br w:type="column"/>。"""
        from lxml import etree
        br = etree.SubElement(run._r, qn('w:br'))
        br.set(qn('w:type'), 'column')

    def _get_cell(self, table_idx: int, row_idx: int, col_idx: int):
        """
        获取指定表格单元格，含边界检查。

        对于含复杂合并单元格（gridSpan / vMerge）的表格，
        python-docx 的 table.cell(r, c) 可能崩溃。
        本方法先尝试标准 API，失败后自动回退到 XML 级别的定位。
        """
        from docx.oxml.ns import qn as _qn
        from docx.table import _Cell

        tables = self._doc.tables
        if table_idx >= len(tables):
            raise IndexError(f"表格索引 {table_idx} 超出范围，共 {len(tables)} 个表格")
        table = tables[table_idx]
        if row_idx >= len(table.rows):
            raise IndexError(f"行索引 {row_idx} 超出范围，表格 {table_idx} 共 {len(table.rows)} 行")

        # 优先尝试标准 API
        try:
            return table.cell(row_idx, col_idx)
        except (IndexError, KeyError):
            pass

        # 回退：直接遍历 XML，用 gridSpan 计算逻辑列号
        tbl = table._tbl
        tr_list = tbl.findall(_qn('w:tr'))
        if row_idx >= len(tr_list):
            raise IndexError(f"行索引 {row_idx} 超出范围（XML 级别）")

        tr = tr_list[row_idx]
        tc_list = tr.findall(_qn('w:tc'))
        logical_col = 0
        for tc in tc_list:
            grid_span = 1
            tcPr = tc.find(_qn('w:tcPr'))
            if tcPr is not None:
                gs_el = tcPr.find(_qn('w:gridSpan'))
                if gs_el is not None:
                    grid_span = int(gs_el.get(_qn('w:val'), '1'))

            if logical_col <= col_idx < logical_col + grid_span:
                return _Cell(tc, table)
            logical_col += grid_span

        raise IndexError(
            f"列索引 {col_idx} 超出范围，表格 {table_idx} 行 {row_idx} 的逻辑列数为 {logical_col}"
        )

    @staticmethod
    def _alignment_str(alignment) -> Optional[str]:
        """将 WD_ALIGN_PARAGRAPH 枚举转为可读字符串。"""
        if alignment is None:
            return None
        mapping = {0: 'left', 1: 'center', 2: 'right', 3: 'justify', 4: 'distribute'}
        return mapping.get(int(alignment), str(alignment))

    @staticmethod
    def _alignment_enum(name: str):
        """将字符串对齐名转换为 WD_ALIGN_PARAGRAPH 枚举。"""
        m = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'distribute': WD_ALIGN_PARAGRAPH.DISTRIBUTE,
        }
        if name not in m:
            raise ValueError(f"未知对齐方式 '{name}'，支持: {list(m.keys())}")
        return m[name]

    @staticmethod
    def _to_length(value):
        """
        将 width/height 参数标准化为 docx Length 对象。
        - None         → None（使用图片原始尺寸）
        - Length 对象  → 原样返回（Inches/Cm/Pt/Emu 等）
        - int / float  → 视为英寸
        """
        if value is None:
            return None
        if isinstance(value, Length):
            return value
        if isinstance(value, (int, float)):
            return Inches(value)
        raise TypeError(
            f"width/height 必须是 None、float（英寸）或 docx Length 对象，"
            f"收到: {type(value).__name__}"
        )
