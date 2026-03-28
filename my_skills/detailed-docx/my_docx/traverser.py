"""
traverser.py
------------
统一文档遍历器。

将 Word 文档中所有「可包含文本」的容器展平为统一的迭代接口：
  - 正文段落 (body paragraphs)
  - 表格单元格中的段落 (table → row → cell → paragraph)
  - 页眉段落 (header paragraphs)
  - 页脚段落 (footer paragraphs)

每个被遍历的段落都携带 LocationInfo 上下文，
告诉调用者它来自哪里（正文第几段、第几个表格第几行第几列等）。
"""

from typing import List, Iterator, Optional
from dataclasses import dataclass, field

from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.table import Table


@dataclass
class LocationInfo:
    """描述段落在文档中所处的位置。"""
    container: str  # 'body' | 'table' | 'header' | 'footer'
    para_idx: int = -1  # 段落在其容器中的索引
    # 以下字段仅 container == 'table' 时使用
    table_idx: int = -1
    row_idx: int = -1
    col_idx: int = -1
    # 以下字段仅 container in ('header', 'footer') 时使用
    section_idx: int = -1

    def to_dict(self) -> dict:
        d = {'container': self.container, 'para_idx': self.para_idx}
        if self.container == 'table':
            d.update(table_idx=self.table_idx, row_idx=self.row_idx, col_idx=self.col_idx)
        if self.container in ('header', 'footer'):
            d['section_idx'] = self.section_idx
        return d

    def __repr__(self) -> str:
        if self.container == 'body':
            return f'body:para[{self.para_idx}]'
        if self.container == 'table':
            return f'table[{self.table_idx}]:row[{self.row_idx}]:col[{self.col_idx}]:para[{self.para_idx}]'
        return f'{self.container}:section[{self.section_idx}]:para[{self.para_idx}]'


@dataclass
class LocatedParagraph:
    """段落 + 位置信息的捆绑。"""
    paragraph: Paragraph
    location: LocationInfo


class DocumentTraverser:
    """
    文档遍历器。

    用法：
        traverser = DocumentTraverser(document)
        for lp in traverser.iter_all():
            print(lp.location, lp.paragraph.text)
    """

    def __init__(self, document: Document):
        self._doc = document

    # ─── 公开遍历方法 ───────────────

    def iter_body(self) -> Iterator[LocatedParagraph]:
        """遍历正文中的所有段落。"""
        for idx, para in enumerate(self._doc.paragraphs):
            yield LocatedParagraph(
                paragraph=para,
                location=LocationInfo(container='body', para_idx=idx)
            )

    def iter_tables(self, table_indices: Optional[List[int]] = None) -> Iterator[LocatedParagraph]:
        """
        遍历表格中的所有段落（含嵌套表格，深度优先）。

        Args:
            table_indices: 若指定，只遍历这些索引的表格。None = 遍历全部。
        """
        for t_idx, table in enumerate(self._doc.tables):
            if table_indices is not None and t_idx not in table_indices:
                continue
            yield from self._iter_table(table, t_idx)

    def _iter_table(self, table: Table, t_idx: int) -> Iterator[LocatedParagraph]:
        """递归遍历一个表格（含嵌套表格）。"""
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    yield LocatedParagraph(
                        paragraph=para,
                        location=LocationInfo(
                            container='table',
                            para_idx=p_idx,
                            table_idx=t_idx,
                            row_idx=r_idx,
                            col_idx=c_idx,
                        )
                    )
                # 处理嵌套表格
                for nested_table in cell.tables:
                    yield from self._iter_table(nested_table, t_idx)

    def iter_headers(self) -> Iterator[LocatedParagraph]:
        """遍历所有节（Section）的页眉段落。"""
        for s_idx, section in enumerate(self._doc.sections):
            header = section.header
            if header.is_linked_to_previous:
                continue
            for p_idx, para in enumerate(header.paragraphs):
                yield LocatedParagraph(
                    paragraph=para,
                    location=LocationInfo(
                        container='header', para_idx=p_idx, section_idx=s_idx
                    )
                )

    def iter_footers(self) -> Iterator[LocatedParagraph]:
        """遍历所有节（Section）的页脚段落。"""
        for s_idx, section in enumerate(self._doc.sections):
            footer = section.footer
            if footer.is_linked_to_previous:
                continue
            for p_idx, para in enumerate(footer.paragraphs):
                yield LocatedParagraph(
                    paragraph=para,
                    location=LocationInfo(
                        container='footer', para_idx=p_idx, section_idx=s_idx
                    )
                )

    def iter_all(self, scope: str = 'all') -> Iterator[LocatedParagraph]:
        """
        按 scope 遍历文档中相应区域的所有段落。

        Args:
            scope:
                'all'     - 正文 + 表格 + 页眉 + 页脚
                'body'    - 仅正文段落
                'tables'  - 仅表格
                'headers' - 仅页眉
                'footers' - 仅页脚
        """
        if scope in ('all', 'body'):
            yield from self.iter_body()
        if scope in ('all', 'tables'):
            yield from self.iter_tables()
        if scope in ('all', 'headers'):
            yield from self.iter_headers()
        if scope in ('all', 'footers'):
            yield from self.iter_footers()
