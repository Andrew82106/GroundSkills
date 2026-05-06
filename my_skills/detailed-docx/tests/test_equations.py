"""
test_equations.py
-----------------
Tests for LaTeX equation insertion (LaTeX → MathML → OMML).
"""
import os
import sys
import pytest
from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from my_docx import DocxEditor


FIXTURE_DIR = os.path.join(os.path.dirname(__file__), 'fixtures')
os.makedirs(FIXTURE_DIR, exist_ok=True)

OMATH_TAG = '{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath'


def _count_omath_in_paragraph(para_xml) -> int:
    return len(list(para_xml.iter(OMATH_TAG)))


def test_add_equation_creates_omath_element():
    editor = DocxEditor.create_new()
    res = editor.add_equation(r"E = mc^2", alignment='center')

    assert res['success'] is True
    assert res['rendered'] is True
    assert res['para_idx'] == 0

    out = os.path.join(FIXTURE_DIR, 'eq_basic.docx')
    editor.save(out)

    # Reopen with raw python-docx and verify <m:oMath> exists
    doc = Document(out)
    para = doc.paragraphs[0]
    assert _count_omath_in_paragraph(para._p) == 1


def test_add_equation_to_paragraph_inline():
    editor = DocxEditor.create_new()
    editor.add_paragraph("根据牛顿第二定律：")
    res = editor.add_equation_to_paragraph(0, r"F = ma")

    assert res['rendered'] is True
    out = os.path.join(FIXTURE_DIR, 'eq_inline.docx')
    editor.save(out)

    doc = Document(out)
    para0 = doc.paragraphs[0]
    assert "根据牛顿第二定律" in para0.text
    assert _count_omath_in_paragraph(para0._p) == 1


def test_add_equation_complex_latex():
    editor = DocxEditor.create_new()
    res = editor.add_equation(r"\frac{-b \pm \sqrt{b^2 - 4ac}}{2a}")
    assert res['rendered'] is True

    out = os.path.join(FIXTURE_DIR, 'eq_quadratic.docx')
    editor.save(out)
    doc = Document(out)
    assert _count_omath_in_paragraph(doc.paragraphs[0]._p) == 1


def test_add_equation_failure_writes_fallback_text():
    editor = DocxEditor.create_new()
    # 故意传入 latex2mathml 解析失败的内容（未闭合的 \frac）
    bad = r"\frac{a"
    res = editor.add_equation(bad)

    assert res['success'] is True
    assert res['rendered'] is False
    assert 'error' in res

    out = os.path.join(FIXTURE_DIR, 'eq_fallback.docx')
    editor.save(out)

    reopened = DocxEditor(out)
    smap = reopened.get_structural_map()
    fallback_text = smap['paragraphs'][0]['text']
    assert '[渲染失败]' in fallback_text
    assert bad in fallback_text

    # 失败时不应留下 oMath 元素
    doc = Document(out)
    assert _count_omath_in_paragraph(doc.paragraphs[0]._p) == 0


def test_equation_does_not_break_subsequent_replace():
    """插入公式后，对其他段落做 replace_text 不应报错。"""
    editor = DocxEditor.create_new()
    editor.add_paragraph("替换前")
    editor.add_equation(r"a + b = c")
    out = os.path.join(FIXTURE_DIR, 'eq_replace_safe.docx')
    editor.save(out)

    e2 = DocxEditor(out)
    n = e2.replace_text("替换前", "替换后")['replaced_count']
    assert n == 1
    e2.save(out)

    e3 = DocxEditor(out)
    doc = Document(out)
    # 公式仍在
    assert _count_omath_in_paragraph(doc.paragraphs[1]._p) == 1
    assert any('替换后' in p['text'] for p in e3.get_structural_map()['paragraphs'])
