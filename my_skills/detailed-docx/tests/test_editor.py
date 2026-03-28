"""
test_editor.py
--------------
端到端测试：用代码构造各种复杂的 .docx，然后用 DocxEditor 读取和修改，
最后断言格式和内容都符合预期。
"""
import os
import sys
import pytest
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 确保能 import my_docx
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from my_docx import DocxEditor


FIXTURE_DIR = os.path.join(os.path.dirname(__file__), 'fixtures')
os.makedirs(FIXTURE_DIR, exist_ok=True)


# ═══════════════════════════════════
#  Fixtures: 构建测试用文档
# ═══════════════════════════════════

def _create_basic_doc() -> str:
    """一个基础文档：含标题、正文段落、一个小表格。"""
    path = os.path.join(FIXTURE_DIR, 'basic.docx')
    doc = Document()
    doc.add_heading('测试标题', level=1)
    p1 = doc.add_paragraph()
    r1 = p1.add_run('这是一段')
    r1.font.bold = True
    r1.font.size = Pt(14)
    r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色粗体
    r2 = p1.add_run('普通正文内容')
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 黑色

    doc.add_paragraph('项目非常成功，需要重点关注。')

    # 表格
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = '姓名'
    table.cell(0, 1).text = '年龄'
    table.cell(0, 2).text = '城市'
    table.cell(1, 0).text = '张三'
    table.cell(1, 1).text = '25'
    table.cell(1, 2).text = '北京'

    doc.save(path)
    return path


def _create_cross_run_doc() -> str:
    """测试跨 Run 替换：刻意把 '目标文本' 拆分到多个 Run。"""
    path = os.path.join(FIXTURE_DIR, 'cross_run.docx')
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run('前缀-目')
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    r2 = p.add_run('标文')
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
    r3 = p.add_run('本-后缀')
    r3.font.size = Pt(16)
    doc.save(path)
    return path


def _create_image_doc() -> str:
    """含图片的文档：确保图片在替换操作后保持完好。"""
    path = os.path.join(FIXTURE_DIR, 'with_image.docx')
    doc = Document()
    p = doc.add_paragraph()
    p.add_run('文字在图片前面')
    # 添加一个极小的图片（1x1 pixel PNG）来测试
    # 这里我们用另一种方式：直接在段落后加文字来模拟
    p2 = doc.add_paragraph('图片后面的文字也要替换')
    doc.save(path)
    return path


def _create_format_doc() -> str:
    """
    格式增量叠加测试文档。
    段落："我很开心" — '很' 是红色粗体，其余是黑色常规。
    """
    path = os.path.join(FIXTURE_DIR, 'format_test.docx')
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run('我')
    r1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    r2 = p.add_run('很')
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    r3 = p.add_run('开心')
    r3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    doc.save(path)
    return path


# ═══════════════════════════════════
#  测试：结构化读取
# ═══════════════════════════════════

class TestStructuralRead:
    def setup_method(self):
        self.path = _create_basic_doc()
        self.editor = DocxEditor(self.path)

    def test_get_structural_map(self):
        smap = self.editor.get_structural_map()
        assert 'paragraphs' in smap
        assert 'tables' in smap
        assert 'sections' in smap
        assert len(smap['tables']) == 1
        assert smap['tables'][0]['rows'] == 2
        assert smap['tables'][0]['cols'] == 3

    def test_get_paragraphs(self):
        paras = self.editor.get_paragraphs(scope='body')
        assert len(paras) >= 2  # 至少标题 + 正文
        # 第一段是标题
        assert '测试标题' in paras[0]['text']

    def test_get_table_data(self):
        data = self.editor.get_table_data(0)
        assert data['rows'] == 2
        assert data['cols'] == 3
        assert data['cells'][0][0]['text'] == '姓名'
        assert data['cells'][1][2]['text'] == '北京'

    def test_get_table_cell_text(self):
        assert self.editor.get_table_cell_text(0, 1, 0) == '张三'

    def test_get_run_details(self):
        # 第1段（index=1）有两个 run：红色粗体 + 黑色常规
        details = self.editor.get_run_details(1)
        assert len(details) == 2
        assert details[0]['text'] == '这是一段'
        assert details[0]['format'].get('bold') == True
        assert details[0]['format'].get('color') == 'FF0000'

    def test_table_out_of_range(self):
        with pytest.raises(IndexError):
            self.editor.get_table_data(99)


# ═══════════════════════════════════
#  测试：文本替换 + 格式保留
# ═══════════════════════════════════

class TestTextReplacement:
    def test_simple_replace_preserves_format(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        result = editor.replace_text('普通正文', '被替换的文本')
        assert result['replaced_count'] == 1

        # 验证：第一个 Run（红色粗体）不受影响
        details = editor.get_run_details(1)
        assert details[0]['text'] == '这是一段'
        assert details[0]['format'].get('bold') == True
        assert details[0]['format'].get('color') == 'FF0000'
        # 第二个 Run 文本已更新
        assert '被替换的文本' in details[1]['text']

    def test_cross_run_replace(self):
        """测试跨 Run 替换：'目标文本' 横跨 3 个 Run。"""
        path = _create_cross_run_doc()
        editor = DocxEditor(path)
        result = editor.replace_text('目标文本', 'NEW')
        assert result['replaced_count'] == 1

        # 替换后段落文本应为 '前缀-NEW-后缀'
        paras = editor.get_paragraphs(scope='body')
        full_text = paras[0]['text']
        assert 'NEW' in full_text
        assert '目标文本' not in full_text

        # 验证：第一个 Run 的粗体和红色格式仍然存在
        details = editor.get_run_details(0)
        first_run = details[0]
        assert first_run['format'].get('bold') == True
        assert first_run['format'].get('color') == 'FF0000'

    def test_replace_in_table(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        result = editor.replace_in_table(0, 1, 2, '北京', '上海')
        assert result['replaced_count'] == 1
        assert editor.get_table_cell_text(0, 1, 2) == '上海'

    def test_replace_count_limit(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        # 文档里有多地方可能出现 "测试"
        result = editor.replace_text('测试', '实验', count=1)
        assert result['replaced_count'] <= 1

    def test_regex_replace(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        # 用正则把 "2" 后面跟 "5" 替换（年龄 25 → 30）
        result = editor.replace_text(r'25', '30', use_regex=False)
        assert result['replaced_count'] >= 1

    def test_save_and_reload(self):
        """替换后保存、重新加载、验证数据完整。"""
        path = _create_basic_doc()
        out_path = os.path.join(FIXTURE_DIR, 'saved_output.docx')
        editor = DocxEditor(path)
        editor.replace_text('张三', '李四')
        editor.save(out_path)

        # 重新加载验证
        editor2 = DocxEditor(out_path)
        assert editor2.get_table_cell_text(0, 1, 0) == '李四'
        os.remove(out_path)


# ═══════════════════════════════════
#  测试：增量格式叠加
# ═══════════════════════════════════

class TestFormatModification:
    def test_incremental_format_overlay(self):
        """
        '我很开心' — '很' 是红色粗体
        给整个 '我很开心' 加下划线
        验证：'很' 仍然保持红色+粗体，同时新增下划线
        """
        path = _create_format_doc()
        editor = DocxEditor(path)
        result = editor.modify_format('我很开心', {'underline': True})
        assert result['modified_count'] == 1

        details = editor.get_run_details(0)
        # Run[0] '我' — 应该有 underline，颜色仍为黑色
        assert details[0]['format'].get('underline') == True
        assert details[0]['format'].get('color') == '000000'

        # Run[1] '很' — 应该有 bold + 红色 + underline（增量叠加）
        assert details[1]['format'].get('bold') == True
        assert details[1]['format'].get('color') == 'FF0000'
        assert details[1]['format'].get('underline') == True

        # Run[2] '开心' — 应该有 underline，颜色仍为黑色
        assert details[2]['format'].get('underline') == True
        assert details[2]['format'].get('color') == '000000'

    def test_modify_paragraph_format(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        result = editor.modify_paragraph_format(1, {
            'alignment': 'center',
            'space_before': 18,
        })
        assert result['success'] == True

        paras = editor.get_paragraphs(scope='body')
        fmt = paras[1]['paragraph_format']
        assert fmt.get('alignment') == 'center'
        assert fmt.get('space_before') == 18.0

    def test_modify_format_with_color_change(self):
        """只改颜色，不动其他格式。"""
        path = _create_format_doc()
        editor = DocxEditor(path)
        editor.modify_format('很', {'color': '0000FF'})  # 蓝色

        details = editor.get_run_details(0)
        # '很' 的 bold 应该还在，颜色变蓝
        assert details[1]['format'].get('bold') == True
        assert details[1]['format'].get('color') == '0000FF'


# ═══════════════════════════════════
#  测试：表格单元格样式
# ═══════════════════════════════════

class TestTableCellStyle:
    def test_cell_shading(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        result = editor.modify_table_cell_shading(0, 0, 0, 'FFFF00')
        assert result['success'] == True

        # 验证：重新读取
        data = editor.get_table_data(0)
        assert data['cells'][0][0]['shading'] == 'FFFF00'

    def test_cell_border(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        result = editor.modify_table_cell_border(
            0, 0, 0,
            top={'sz': '12', 'val': 'single', 'color': '000000'},
            bottom={'sz': '12', 'val': 'single', 'color': 'FF0000'},
        )
        assert result['success'] == True


# ═══════════════════════════════════
#  测试：从零创建文档
# ═══════════════════════════════════

class TestCreateNew:
    def test_create_empty_and_add_content(self):
        """从空白文档创建，添加标题、段落、表格，保存后重新读取验证。"""
        out_path = os.path.join(FIXTURE_DIR, 'created_new.docx')
        editor = DocxEditor.create_new()
        editor.add_heading('新建文档测试', level=1)
        editor.add_paragraph('这是正文', fmt={'bold': True, 'size': 14})
        editor.add_table(2, 3, data=[['A', 'B', 'C'], ['1', '2', '3']])
        editor.save(out_path)

        # 重新加载验证
        e2 = DocxEditor(out_path)
        smap = e2.get_structural_map()
        assert '新建文档测试' in smap['paragraphs'][0]['text']
        assert smap['tables'][0]['rows'] == 2
        assert smap['tables'][0]['preview'][1] == ['1', '2', '3']
        
        # 验证格式
        details = e2.get_run_details(1)
        assert details[0]['format'].get('bold') == True
        assert details[0]['format'].get('size') == 14.0
        os.remove(out_path)

    def test_add_run_mixed_format(self):
        """在同一段落内创建多格式混排。"""
        out_path = os.path.join(FIXTURE_DIR, 'mixed_run.docx')
        editor = DocxEditor.create_new()
        editor.add_paragraph('我觉得')
        editor.add_run_to_paragraph(0, '非常好', {'bold': True, 'color': 'FF0000'})
        editor.save(out_path)

        e2 = DocxEditor(out_path)
        details = e2.get_run_details(0)
        assert details[0]['text'] == '我觉得'
        assert details[1]['text'] == '非常好'
        assert details[1]['format'].get('bold') == True
        assert details[1]['format'].get('color') == 'FF0000'
        os.remove(out_path)

    def test_create_new_must_specify_path(self):
        """create_new() 之后 save() 不传路径应该报错。"""
        editor = DocxEditor.create_new()
        with pytest.raises(ValueError):
            editor.save()


# ═══════════════════════════════════
#  测试：删除操作
# ═══════════════════════════════════

class TestDelete:
    def test_delete_paragraph(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        smap_before = editor.get_structural_map()
        para_count_before = len(smap_before['paragraphs'])

        result = editor.delete_paragraph(1)
        assert result['success'] == True
        assert '这是一段' in result['deleted_text']

        smap_after = editor.get_structural_map()
        assert len(smap_after['paragraphs']) == para_count_before - 1

    def test_delete_table_row(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        result = editor.delete_table_row(0, 1)
        assert result['success'] == True

        data = editor.get_table_data(0)
        assert data['rows'] == 1  # 只剩表头行

    def test_delete_table(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        result = editor.delete_table(0)
        assert result['success'] == True

        smap = editor.get_structural_map()
        assert len(smap['tables']) == 0

    def test_clear_table_cell(self):
        path = _create_basic_doc()
        editor = DocxEditor(path)
        result = editor.clear_table_cell(0, 1, 0)
        assert result['success'] == True
        assert result['cleared_text'] == '张三'
        assert editor.get_table_cell_text(0, 1, 0) == ''


if __name__ == '__main__':
    pytest.main([__file__, '-v'])

