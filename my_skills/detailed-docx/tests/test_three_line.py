"""
test_three_line.py
------------------
三线表（booktabs）API 的测试。
"""

import os
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from my_docx import DocxEditor
from docx import Document
from docx.oxml.ns import qn


@pytest.fixture
def tmp_path():
    with tempfile.TemporaryDirectory() as d:
        yield d


class TestThreeLineStyle:
    """apply_three_line_style 基本功能。"""

    def test_basic_three_line(self, tmp_path):
        """基本三线表：第一行表头。"""
        editor = DocxEditor.create_new()
        editor.add_table(4, 3, data=[
            ["方法", "精确率", "召回率"],
            ["A", "90.1", "85.3"],
            ["B", "88.5", "92.1"],
            ["C", "91.2", "87.6"],
        ])
        result = editor.apply_three_line_style(0)
        assert result['success'] is True
        assert result['header_rows'] == 1

        path = os.path.join(tmp_path, 'three_line.docx')
        editor.save(path)

        # 验证：重新打开，检查边框
        doc = Document(path)
        table = doc.tables[0]

        # 第一行应有 top 边框
        cell00 = table.rows[0].cells[0]
        tc = cell00._tc
        tcPr = tc.find(qn('w:tcPr'))
        borders = tcPr.find(qn('w:tcBorders')) if tcPr is not None else None
        assert borders is not None
        top = borders.find(qn('w:top'))
        assert top is not None
        assert top.get(qn('w:val')) == 'single'

        # 最后一行应有 bottom 边框
        last_cell = table.rows[-1].cells[0]
        tc_last = last_cell._tc
        tcPr_last = tc_last.find(qn('w:tcPr'))
        borders_last = tcPr_last.find(qn('w:tcBorders')) if tcPr_last is not None else None
        assert borders_last is not None
        bottom = borders_last.find(qn('w:bottom'))
        assert bottom is not None
        assert bottom.get(qn('w:val')) == 'single'

    def test_two_header_rows(self, tmp_path):
        """双行表头的三线表。"""
        editor = DocxEditor.create_new()
        editor.add_table(5, 3, data=[
            ["类别", "指标A", "指标B"],
            ["", "子指标1", "子指标2"],
            ["方法1", "90", "85"],
            ["方法2", "88", "92"],
            ["方法3", "91", "87"],
        ])
        result = editor.apply_three_line_style(0, header_rows=2)
        assert result['success'] is True
        assert result['header_rows'] == 2

        path = os.path.join(tmp_path, 'two_header.docx')
        editor.save(path)

        # 栏目线应在第2行(index=1)的 bottom
        doc = Document(path)
        cell_h = doc.tables[0].rows[1].cells[0]
        tc = cell_h._tc
        tcPr = tc.find(qn('w:tcPr'))
        borders = tcPr.find(qn('w:tcBorders'))
        assert borders is not None
        bottom = borders.find(qn('w:bottom'))
        assert bottom is not None

    def test_no_vertical_borders(self, tmp_path):
        """三线表不应有竖线。"""
        editor = DocxEditor.create_new()
        editor.add_table(3, 3, data=[
            ["A", "B", "C"],
            ["1", "2", "3"],
            ["4", "5", "6"],
        ], style='Table Grid')  # Table Grid 默认有竖线
        editor.apply_three_line_style(0)

        path = os.path.join(tmp_path, 'no_vert.docx')
        editor.save(path)

        doc = Document(path)
        tbl = doc.tables[0]._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        tblBorders = tblPr.find(qn('w:tblBorders'))
        assert tblBorders is not None

        # insideV 和 left/right 应该是 nil
        insideV = tblBorders.find(qn('w:insideV'))
        assert insideV is not None
        assert insideV.get(qn('w:val')) == 'nil'

        left = tblBorders.find(qn('w:left'))
        assert left is not None
        assert left.get(qn('w:val')) == 'nil'

    def test_custom_border_sizes(self, tmp_path):
        """自定义线条粗细。"""
        editor = DocxEditor.create_new()
        editor.add_table(3, 2, data=[["H1", "H2"], ["a", "b"], ["c", "d"]])
        editor.apply_three_line_style(0, top_border_pt=2.0, mid_border_pt=1.0, bottom_border_pt=2.0)

        path = os.path.join(tmp_path, 'custom.docx')
        editor.save(path)

        doc = Document(path)
        cell = doc.tables[0].rows[0].cells[0]
        tc = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        borders = tcPr.find(qn('w:tcBorders'))
        top = borders.find(qn('w:top'))
        # 2.0pt * 8 = 16
        assert top.get(qn('w:sz')) == '16'

    def test_index_out_of_range(self):
        """表格索引越界。"""
        editor = DocxEditor.create_new()
        editor.add_table(2, 2)
        with pytest.raises(IndexError):
            editor.apply_three_line_style(99)

    def test_header_rows_too_large(self):
        """header_rows >= 总行数应报错。"""
        editor = DocxEditor.create_new()
        editor.add_table(3, 2)
        with pytest.raises(ValueError, match="header_rows"):
            editor.apply_three_line_style(0, header_rows=3)

    def test_body_rows_no_borders(self, tmp_path):
        """表体中间行不应有任何边框。"""
        editor = DocxEditor.create_new()
        editor.add_table(5, 2, data=[
            ["H1", "H2"],
            ["a", "b"],
            ["c", "d"],
            ["e", "f"],
            ["g", "h"],
        ])
        editor.apply_three_line_style(0)

        path = os.path.join(tmp_path, 'body.docx')
        editor.save(path)

        doc = Document(path)
        # 第2行(index=1)和第3行(index=2)不应有任何 tcBorders
        for row_idx in [1, 2, 3]:
            cell = doc.tables[0].rows[row_idx].cells[0]
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                borders = tcPr.find(qn('w:tcBorders'))
                assert borders is None, f"行{row_idx} 不应有边框"
