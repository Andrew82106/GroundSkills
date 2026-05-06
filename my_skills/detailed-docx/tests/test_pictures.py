"""
test_pictures.py
----------------
Tests for image insertion APIs: add_picture, add_picture_to_paragraph,
add_picture_to_table_cell, count_pictures.
"""
import os
import sys
import struct
import zlib
import pytest
from docx.shared import Inches, Cm

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from my_docx import DocxEditor


FIXTURE_DIR = os.path.join(os.path.dirname(__file__), 'fixtures')
os.makedirs(FIXTURE_DIR, exist_ok=True)


def _make_png(path: str, width: int = 32, height: int = 32) -> str:
    """Create a tiny solid-color PNG without external deps."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack('>I', len(data))
            + tag + data
            + struct.pack('>I', zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    sig = b'\x89PNG\r\n\x1a\n'
    ihdr = struct.pack('>IIBBBBB', width, height, 8, 2, 0, 0, 0)
    raw = b''
    for _ in range(height):
        raw += b'\x00' + (b'\x33\x99\xCC' * width)
    idat = zlib.compress(raw)
    blob = sig + chunk(b'IHDR', ihdr) + chunk(b'IDAT', idat) + chunk(b'IEND', b'')
    with open(path, 'wb') as f:
        f.write(blob)
    return path


@pytest.fixture(scope='module')
def png_path():
    return _make_png(os.path.join(FIXTURE_DIR, 'tiny.png'))


def test_add_picture_appends_paragraph(png_path):
    editor = DocxEditor.create_new()
    editor.add_paragraph('正文内容')
    res = editor.add_picture(png_path, width=1.5, alignment='center')

    assert res['success'] is True
    assert res['para_idx'] == 1
    out = os.path.join(FIXTURE_DIR, 'pic_basic.docx')
    editor.save(out)

    reopened = DocxEditor(out)
    assert reopened.count_pictures()['count'] == 1
    smap = reopened.get_structural_map()
    assert smap['paragraphs'][1]['alignment'] == 'center'


def test_add_picture_to_paragraph_appends_run(png_path):
    editor = DocxEditor.create_new()
    editor.add_paragraph('图片在我后面：')
    res = editor.add_picture_to_paragraph(0, png_path, width=Cm(1))

    assert res['success'] is True
    assert res['para_idx'] == 0
    assert res['run_idx'] >= 1

    out = os.path.join(FIXTURE_DIR, 'pic_inline.docx')
    editor.save(out)

    reopened = DocxEditor(out)
    assert reopened.count_pictures()['count'] == 1
    runs = reopened.get_run_details(0)
    assert any(not r['is_text_run'] for r in runs), '图片 Run 应被识别为非文本 Run'


def test_add_picture_to_table_cell(png_path):
    editor = DocxEditor.create_new()
    editor.add_table(2, 2, data=[['表头', '图'], ['A', 'B']])
    res = editor.add_picture_to_table_cell(
        0, 0, 1, png_path, width=Inches(0.5), alignment='center', clear_cell=True,
    )
    assert res['success'] is True

    out = os.path.join(FIXTURE_DIR, 'pic_table.docx')
    editor.save(out)

    reopened = DocxEditor(out)
    assert reopened.count_pictures()['count'] == 1
    cell_text = reopened.get_table_cell_text(0, 0, 1)
    assert cell_text == '', 'clear_cell=True 应清掉 "图" 文字'


def test_text_replace_does_not_break_picture(png_path):
    """Image should survive a replace_text call in the same document."""
    editor = DocxEditor.create_new()
    editor.add_paragraph('替换我')
    editor.add_picture(png_path, width=1.0)
    out = os.path.join(FIXTURE_DIR, 'pic_replace_safe.docx')
    editor.save(out)

    e2 = DocxEditor(out)
    e2.replace_text('替换我', '已替换')
    e2.save(out)

    e3 = DocxEditor(out)
    assert e3.count_pictures()['count'] == 1
    assert any('已替换' in p['text'] for p in e3.get_structural_map()['paragraphs'])


def test_to_length_rejects_bad_type():
    editor = DocxEditor.create_new()
    with pytest.raises(TypeError):
        editor._to_length('3in')
