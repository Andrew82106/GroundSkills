#!/usr/bin/env python3
"""
模板应用器 - 核心功能
分析模板文档的格式，应用到目标文档
用法：python apply_template.py 模板.docx 目标.docx 输出.docx
"""

import sys
import copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


def get_paragraph_style(para):
    """提取段落样式"""
    style = {}
    pf = para.paragraph_format

    # 对齐方式
    if pf.alignment is not None:
        style['alignment'] = pf.alignment

    # 首行缩进
    if pf.first_line_indent is not None:
        style['first_line_indent'] = pf.first_line_indent

    # 左缩进
    if pf.left_indent is not None:
        style['left_indent'] = pf.left_indent

    # 右缩进
    if pf.right_indent is not None:
        style['right_indent'] = pf.right_indent

    # 行距
    if pf.line_spacing is not None:
        style['line_spacing'] = pf.line_spacing
    if pf.line_spacing_rule is not None:
        style['line_spacing_rule'] = pf.line_spacing_rule

    # 段前段后距
    if pf.space_before is not None:
        style['space_before'] = pf.space_before
    if pf.space_after is not None:
        style['space_after'] = pf.space_after

    return style


def get_run_style(run):
    """提取 run 样式（字体、字号等）"""
    style = {}
    font = run.font

    if font.name is not None:
        style['font_name'] = font.name
    if font.size is not None:
        style['font_size'] = font.size
    if font.bold is not None:
        style['bold'] = font.bold
    if font.italic is not None:
        style['italic'] = font.italic
    if font.underline is not None:
        style['underline'] = font.underline

    # 中文字体
    if hasattr(font, 'eastAsia') and font.eastAsia is not None:
        style['eastAsia'] = font.eastAsia

    return style


def analyze_template(template_path):
    """分析模板文档，提取样式"""
    print(f"分析模板：{template_path}")
    doc = Document(template_path)

    template_styles = {
        'page': {},
        'title': {},      # 主标题样式
        'heading1': {},   # 一级标题
        'heading2': {},   # 二级标题
        'body': {},       # 正文
        'paragraph': {}   # 段落
    }

    # 提取页面设置（从第一个 section）
    if doc.sections:
        section = doc.sections[0]
        template_styles['page'] = {
            'page_width': section.page_width,
            'page_height': section.page_height,
            'top_margin': section.top_margin,
            'bottom_margin': section.bottom_margin,
            'left_margin': section.left_margin,
            'right_margin': section.right_margin,
        }

    # 遍历段落，识别不同样式
    title_found = False
    heading1_count = 0
    heading2_count = 0
    body_samples = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检测标题类型
        is_title = False
        is_heading1 = False
        is_heading2 = False

        # 检查段落样式名称
        style_name = para.style.name if para.style else ""

        if 'Title' in style_name or (len(text) < 50 and para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER and body_samples == 0):
            is_title = True
        elif 'Heading 1' in style_name or 'Header' in style_name or (len(text) < 30 and para.runs and para.runs[0].font.bold and heading1_count < 3):
            is_heading1 = True
        elif 'Heading 2' in style_name or (len(text) < 30 and para.runs and para.runs[0].font.bold and heading2_count < 5):
            is_heading2 = True

        # 提取样式
        para_style = get_paragraph_style(para)
        run_style = get_run_style(para.runs[0]) if para.runs else {}
        combined_style = {**para_style, **run_style}

        if is_title and not title_found:
            template_styles['title'] = combined_style
            title_found = True
            print(f"  识别主标题样式：{combined_style.get('font_name', '未知')} {combined_style.get('font_size', '未知')}")
        elif is_heading1 and heading1_count < 3:
            template_styles['heading1'] = combined_style
            heading1_count += 1
        elif is_heading2 and heading2_count < 3:
            template_styles['heading2'] = combined_style
            heading2_count += 1
        elif body_samples < 5:
            # 合并正文字体样式
            if not template_styles['body']:
                template_styles['body'] = combined_style
            else:
                # 合并多个样本
                for k, v in combined_style.items():
                    if k not in template_styles['body']:
                        template_styles['body'][k] = v
            body_samples += 1

    # 默认正文样式（如果没检测到）
    if not template_styles['body']:
        template_styles['body'] = {
            'font_name': '仿宋',
            'font_size': Pt(16),
            'first_line_indent': Pt(28),  # 2 字符
            'line_spacing_rule': WD_LINE_SPACING.EXACTLY,
            'line_spacing': Pt(28),
        }

    return template_styles


def apply_style_to_paragraph(para, style, style_type='body'):
    """应用样式到段落"""
    pf = para.paragraph_format

    # 应用段落样式
    if 'alignment' in style:
        pf.alignment = style['alignment']
    if 'first_line_indent' in style:
        pf.first_line_indent = style['first_line_indent']
    if 'left_indent' in style:
        pf.left_indent = style['left_indent']
    if 'right_indent' in style:
        pf.right_indent = style['right_indent']
    if 'line_spacing' in style:
        pf.line_spacing = style['line_spacing']
    if 'line_spacing_rule' in style:
        pf.line_spacing_rule = style['line_spacing_rule']
    if 'space_before' in style:
        pf.space_before = style['space_before']
    if 'space_after' in style:
        pf.space_after = style['space_after']

    # 应用 run 样式（字体、字号）
    for run in para.runs:
        if 'font_name' in style:
            run.font.name = style['font_name']
        if 'eastAsia' in style:
            run.font.eastAsia = style['eastAsia']
        if 'font_size' in style:
            run.font.size = style['font_size']
        # 保留原有的粗体/斜体


def apply_page_setup(doc, page_style):
    """应用页面设置"""
    for section in doc.sections:
        if 'page_width' in page_style:
            section.page_width = page_style['page_width']
        if 'page_height' in page_style:
            section.page_height = page_style['page_height']
        if 'top_margin' in page_style:
            section.top_margin = page_style['top_margin']
        if 'bottom_margin' in page_style:
            section.bottom_margin = page_style['bottom_margin']
        if 'left_margin' in page_style:
            section.left_margin = page_style['left_margin']
        if 'right_margin' in page_style:
            section.right_margin = page_style['right_margin']


def detect_paragraph_type(para, prev_types):
    """检测段落类型（标题、正文等）"""
    text = para.text.strip()
    if not text:
        return 'empty'

    # 检查样式名称
    style_name = para.style.name if para.style else ""

    if 'Title' in style_name:
        return 'title'
    elif 'Heading 1' in style_name:
        return 'heading1'
    elif 'Heading 2' in style_name:
        return 'heading2'

    # 基于内容特征判断
    if len(text) < 40:
        # 短段落，可能是标题
        if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return 'title'
        if para.runs and para.runs[0].font.bold:
            # 检查是否有序号特征
            if any(text.startswith(p) for p in ['一、', '二、', '三、', '1.', '2.', '3.']):
                return 'heading1'
            if any(text.startswith(p) for p in ['（一）', '（二）']):
                return 'heading2'

    return 'body'


def apply_template(template_path, source_path, output_path):
    """将模板样式应用到源文档"""
    # 分析模板
    template_styles = analyze_template(template_path)

    print(f"\n应用样式到：{source_path}")
    source = Document(source_path)

    # 应用页面设置
    apply_page_setup(source, template_styles['page'])
    print(f"  页面设置：{template_styles['page'].get('page_width', 'N/A')} x {template_styles['page'].get('page_height', 'N/A')}")

    # 应用样式到每个段落
    prev_types = []
    body_count = 0
    heading1_count = 0
    heading2_count = 0

    for i, para in enumerate(source.paragraphs):
        if not para.text.strip():
            continue

        para_type = detect_paragraph_type(para, prev_types)
        prev_types.append(para_type)

        if para_type == 'title':
            apply_style_to_paragraph(para, template_styles['title'], 'title')
        elif para_type == 'heading1':
            apply_style_to_paragraph(para, template_styles['heading1'], 'heading1')
            heading1_count += 1
        elif para_type == 'heading2':
            apply_style_to_paragraph(para, template_styles['heading2'], 'heading2')
            heading2_count += 1
        elif para_type == 'body':
            apply_style_to_paragraph(para, template_styles['body'], 'body')
            body_count += 1

    print(f"  处理完成：{body_count} 段正文，{heading1_count} 个一级标题，{heading2_count} 个二级标题")

    # 保存
    source.save(output_path)
    print(f"\n已保存：{output_path}")


def main():
    if len(sys.argv) < 4:
        print("用法：python apply_template.py 模板.docx 目标.docx 输出.docx")
        print()
        print("示例:")
        print("  python apply_template.py template.docx draft.docx output.docx")
        sys.exit(1)

    template_path = sys.argv[1]
    source_path = sys.argv[2]
    output_path = sys.argv[3]

    apply_template(template_path, source_path, output_path)


if __name__ == '__main__':
    main()
